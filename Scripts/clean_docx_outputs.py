import re
from docx import Document
from print_big_dataframe import print_dataframe
from print_big_list import print_big_list, split_list_into_lines
from print_big_text import print_big_text

TRUNCATE_LIMIT = 65
FILE_PATH = 'docs/Introduction-to-`pyspark`.docx'

f = open(FILE_PATH, 'rb')
document = Document(f)
f.close()


pars = document.paragraphs
source_codes = list()
for i in range(len(pars)):
    current_par = pars[i]
    if current_par.style.style_id == 'SourceCode':
        source_codes.append(i)



def par_is_chunk_output(paragraph):
    paragraph_contents = paragraph.iter_inner_content()
    is_chunk_output = False
    for run in paragraph_contents:
        if run.style.style_id == 'VerbatimChar':
            is_chunk_output = True
            break

    return is_chunk_output


def is_stage_output(text):
    pattern = re.compile("^\\[Stage ")
    return pattern.match(text)

chunk_outputs = list()
for index in source_codes:
    current_par = pars[index]
    if par_is_chunk_output(current_par):
        if is_stage_output(current_par.text):
            print(f"[INFO]: Stage output found at index {i}, ignoring output...")
            continue

        chunk_outputs.append(index)
        



def is_dataframe_output(text):
    return text.startswith('+--')

def is_list_output(text):
    regex = re.compile(r'(^\[)|(^StructType\()')
    return regex.match(text)


def need_adjustment(text, n_chars = TRUNCATE_LIMIT):
    lines = text.split('\n')
    for line in lines:
        if len(line) >= n_chars:
            return True
        
    return False


def adjust_chunk_output(text_to_adjust, n_chars = TRUNCATE_LIMIT):
    text_adjusted = text_to_adjust
    if is_dataframe_output(text_to_adjust) and need_adjustment(text_to_adjust, n_chars):
        print("[INFO]: Found a DataFrame output that needs adjustment! Adjusting...")
        return print_dataframe(text_to_adjust, n_chars)
    if is_list_output(text_to_adjust) and need_adjustment(text_to_adjust, n_chars):
        print("[INFO]: Found a list output that needs adjustment! Adjusting...")
        return print_big_list(text_to_adjust, n_chars)
    if need_adjustment(text_to_adjust, n_chars):
        print("[INFO]: Found text output that needs adjustment! Adjusting...")
        return print_big_text(text_to_adjust, n_chars)


    return text_adjusted




(document.styles['VerbatimChar']).font.name = 'Consolas'

for index in chunk_outputs:
    output = (pars[index]).text
    (document.paragraphs[index]).style.font.name = 'Consolas'
    
    adjusted_output = adjust_chunk_output(output)
    (document.paragraphs[index]).text = adjusted_output


document.save("docs/docx_adjusted.docx")

